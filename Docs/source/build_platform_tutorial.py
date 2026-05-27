from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "Docs" / "信息采集板卡固件平台化开发教程.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "16364F"
INK = "243746"
MUTED = "607080"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "E8F3ED"
YELLOW = "FFF4D6"
RED = "FBEAEA"
BORDER = "CBD5E1"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=None, color=None, bold=None, italic=None, name="Microsoft YaHei"):
    run.font.name = name
    props = run._element.get_or_add_rPr()
    fonts = props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        props.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Calibri" if name == "Microsoft YaHei" else name)
    fonts.set(qn("w:hAnsi"), "Calibri" if name == "Microsoft YaHei" else name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=None, bold=None, italic=None, name="Microsoft YaHei"):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def paragraph_border_bottom(paragraph, color=BORDER, size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_box(paragraph, color=BORDER):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), color)
        borders.append(border)
    p_pr.append(borders)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5, color=INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, 29, color=NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.1

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 13, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.line_spacing = 1.2

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        set_style_font(style, 10.5, color=INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = doc.styles.add_style("Code Block", 1)
    set_style_font(code, 8.5, color="233042", name="Consolas")
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.08

    caption = doc.styles["Caption"]
    set_style_font(caption, 9, color=MUTED, italic=True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("DAQ CARD  |  FIRMWARE PLATFORMIZATION GUIDE")
    set_font(r, 8.5, MUTED, bold=True)
    paragraph_border_bottom(header, BORDER, "6", "5")

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("信息采集板卡固件平台化开发教程  |  ")
    set_font(r, 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    for abstract_id, num_id, fmt, text, left, hanging in (
        (42, 42, "bullet", "•", 540, 270),
        (43, 43, "decimal", "%1.", 540, 270),
    ):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        ppr.append(tabs)
        ppr.append(ind)
        level.extend([start, num_fmt, lvl_text, jc, ppr])
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        link = OxmlElement("w:abstractNumId")
        link.set(qn("w:val"), str(abstract_id))
        num.append(link)
        numbering.append(num)


def p(doc, text="", style=None, bold_prefix=None):
    para = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = para.add_run(bold_prefix)
        set_font(first, bold=True)
        para.add_run(text[len(bold_prefix):])
    else:
        para.add_run(text)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para._p.get_or_add_pPr().append(_num_pr(42))
    para.add_run(text)
    return para


def number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para._p.get_or_add_pPr().append(_num_pr(43))
    para.add_run(text)
    return para


def _num_pr(num_id):
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    return num_pr


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.autofit = False
    tbl.allow_autofit = False
    tbl_pr = tbl._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = tbl._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    all_rows = [headers] + rows
    for i, row_data in enumerate(all_rows):
        row = tbl.rows[0] if i == 0 else tbl.add_row()
        if i == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Inches(widths[j] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[j]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                shade(cell, header_fill)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.15
            run = para.add_run(str(text))
            set_font(run, 9.2, color=INK, bold=(i == 0))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return tbl


def callout(doc, label, text, fill=CALLOUT):
    tbl = table(doc, [label], [[text]], [PAGE_WIDTH_DXA], header_fill=fill)
    for cell in tbl.row_cells(0):
        shade(cell, fill)
    for row in tbl.rows:
        for cell in row.cells:
            shade(cell, fill)
    return tbl


def code_block(doc, code, caption_text=None):
    if caption_text:
        para = doc.add_paragraph(caption_text, style="Caption")
        para.paragraph_format.keep_with_next = True
    para = doc.add_paragraph(style="Code Block")
    para.paragraph_format.left_indent = Pt(8)
    para.paragraph_format.right_indent = Pt(8)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(7)
    shade_paragraph(para, "F5F7FA")
    paragraph_box(para)
    lines = code.strip("\n").splitlines()
    for index, line in enumerate(lines):
        run = para.add_run(line)
        set_font(run, size=8.3, color="233042", name="Consolas")
        if index != len(lines) - 1:
            run.add_break()


def cover(doc):
    doc.add_paragraph("FIRMWARE ENGINEERING GUIDE", style="Subtitle").runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)
    title = doc.add_paragraph("信息采集板卡固件\n平台化开发教程", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph("以 STM32F407 + FreeRTOS + CubeMX/Keil 工程为例", style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, BLUE, "14", "12")
    callout(doc, "本文目标", "建立一套可以长期演进的固件分层方法：硬件变化集中在 BSP 与生成层，采集策略、数据协议和产品行为保持稳定可复用。", LIGHT_BLUE)
    meta = [
        ["适用项目", "多通道信息采集板、数据采集终端、传感器网关"],
        ["当前示例工程", "DAQ_card / STM32F407VET6 / FreeRTOS / USART1"],
        ["推荐读者", "嵌入式开发者、固件负责人、板卡联调人员"],
        ["文档版本", "V1.0  |  2026-05-27"],
    ]
    table(doc, ["项目", "说明"], meta, [2100, 7260])
    doc.add_page_break()


def content_pages(doc):
    heading(doc, "阅读路线", 1)
    p(doc, "平台化不是把代码分到更多文件夹，而是建立稳定边界：业务只依赖能力接口，板级代码才知道具体芯片外设和引脚。本教程从你的 DAQ_card 工程出发，先解释思想，再给出能落地的代码组织和推进路线。")
    table(doc, ["章节", "你将获得什么"], [
        ["1-3", "理解平台化目标、当前工程基线和总体架构"],
        ["4-7", "掌握分层职责、数据流、中断与 API 设计"],
        ["8-10", "掌握采集、协议、参数、故障等功能的正确归属"],
        ["11-13", "建立 CubeMX/Keil 工作法、测试策略和迁移路线"],
        ["附录", "目录模板、接口模板、评审清单与术语"],
    ], [1400, 7960])

    heading(doc, "1. 为什么信息采集板需要平台化", 1)
    p(doc, "信息采集固件最早往往从一次硬件点亮开始：ADC 能出数、串口能打印、DMA 能触发，功能就被快速堆在回调或任务里。样机阶段这样很快；一旦出现新板卡、新采样器件、新传输接口、标定和多种工作模式，原来“能跑”的代码会迅速变成修改困难且无法验证的系统。")
    heading(doc, "1.1 平台化要解决的四个变化", 2)
    table(doc, ["变化来源", "未分层时的影响", "平台化后的控制点"], [
        ["换 MCU 或换引脚", "业务文件中到处修改 HAL 句柄与 GPIO", "只修改 Core/CubeMX 与 BSP"],
        ["换 ADC 或采样方式", "协议、任务和硬件控制相互牵连", "BSP 提供统一采样能力，Services 复用"],
        ["UART 改 USB/网口", "采集算法中夹杂发送实现", "协议服务不变，仅替换 transport 适配"],
        ["新增模式与诊断", "主循环不断膨胀、不可测试", "Application 组织状态，Services 分责"],
    ], [1800, 3520, 4040])
    callout(doc, "一句话定义", "平台化固件 = 以稳定接口屏蔽硬件差异，以服务模块沉淀产品能力，以应用层组织设备行为。", GREEN)

    heading(doc, "1.2 平台化成功的判据", 2)
    bullet(doc, "增加一个新的采样通道时，修改范围清晰且可预测。")
    bullet(doc, "将串口传输替换为 USB 或以太网时，采样服务和数据帧编码不需要重写。")
    bullet(doc, "CubeMX 重新生成代码后，手写业务代码不会被覆盖。")
    bullet(doc, "每个模块能说明它的输入、输出、状态和错误，而不是靠全局变量串联。")
    bullet(doc, "硬件尚未完成时，协议、缓冲和状态机仍可以单独测试。")

    heading(doc, "2. 当前 DAQ_card 工程基线", 1)
    p(doc, "本工程已经完成第一次平台骨架整理，适合作为后续采集固件的起点。当前不是成品 DAQ 固件：还没有 ADC、DMA 采集链和数据协议，已有的是可编译的层次与启动入口。")
    table(doc, ["项目状态", "现有事实", "下一步落点"], [
        ["芯片与基础设施", "STM32F407VET6；FreeRTOS；Keil；CubeMX", "保留生成层边界"],
        ["通信", "USART1 初始化，BSP_CONSOLE 已封装写接口", "演进为通用 transport"],
        ["应用入口", "默认任务调用 APP_Init / APP_RunOnce", "增加状态与任务协作"],
        ["采集功能", "DAQ_SERVICE 当前为空闲状态", "接入 ADC + TIM + DMA"],
    ], [1900, 4050, 3410])
    heading(doc, "2.1 现有启动调用链", 2)
    code_block(doc, """
main()
  -> HAL / Clock / GPIO / USART1 initialization     [Core]
  -> MX_FREERTOS_Init() and osKernelStart()         [Core]
     -> StartDefaultTask()                          [Core integration point]
        -> APP_Init() / APP_RunOnce()               [Application]
           -> PLATFORM_Init() / PLATFORM_Process()  [Platform]
              -> BSP_CONSOLE + DAQ_SERVICE          [BSP / Services]
""", "图 1  当前工程的最小平台启动链")
    callout(doc, "重要限制", "当前 `DAQ_SERVICE` 仅提供骨架。不要把它误认为已完成采样；真正采集能力必须在配置 ADC、触发定时器、DMA、缓冲与数据出口之后建立。", YELLOW)

    heading(doc, "3. 总体架构：从生成代码到产品行为", 1)
    p(doc, "一个可维护的信息采集固件应将“设备行为”“通用能力”“板级实现”“生成基础设施”分开。下面的五层结构是本工程后续开发的主干。")
    code_block(doc, """
Application    设备模式、任务协作、命令决策、对外工作流程
     |
Platform       初始化顺序、依赖装配、事件桥接、产品配置
     |
Services       采集引擎、缓冲、协议、参数、诊断、校准
     |
BSP            ADC/DMA、触发时钟、Transport、LED、电源、存储
     |
Core/HAL       CubeMX 初始化、HAL 句柄、中断向量、RTOS glue
""", "图 2  推荐依赖方向")
    heading(doc, "3.1 依赖规则", 2)
    number(doc, "`Application` 可以调用 `Services` 或通过 `Platform` 使用它们，但不能直接操作 HAL 句柄。")
    number(doc, "`Services` 只依赖抽象能力或事件，不知道 ADC 属于哪个外设实例。")
    number(doc, "`BSP` 可以使用 CubeMX 生成的句柄与 HAL API，但不决定产品采集模式。")
    number(doc, "`Core` 只保留生成初始化和极薄的事件转发，不承担业务处理。")
    number(doc, "`Platform` 是装配层，负责把某块板上的 BSP 能力交给通用服务使用。")
    heading(doc, "3.2 什么不叫平台化", 2)
    table(doc, ["表象", "为什么不够", "正确改法"], [
        ["仅新增很多目录", "代码依赖仍混在一起", "用接口与依赖方向约束调用"],
        ["把所有驱动都叫 BSP", "协议与策略仍侵入硬件层", "BSP 只表达板级能力"],
        ["用一个超大 app.c 统一调用", "产品变化仍集中爆炸", "拆成采集、通信、监管任务"],
        ["在中断中直接组帧发送", "实时性与数据一致性不可控", "中断只通知，任务中处理"],
    ], [1800, 3600, 3960])

    heading(doc, "4. 每一层究竟放什么代码", 1)
    heading(doc, "4.1 Core：生成层和最窄接入点", 2)
    p(doc, "`Core/` 是 CubeMX 管理区域，负责时钟树、引脚复用、外设初始化、中断入口和 RTOS 创建。原则是尽量不让手写业务沉入这里，因为重新生成代码时这里最容易变化。")
    table(doc, ["允许放入", "禁止长期堆积"], [
        ["`MX_ADC1_Init()`、`MX_DMA_Init()`、`MX_TIMx_Init()`、IRQ Handler", "采样模式状态机、通信协议、滤波算法"],
        ["`USER CODE` 区中的极短回调转发", "DMA 完成后直接循环发整块数据"],
        ["HAL Tick 与 RTOS 启动 glue", "参数存储、错误策略、设备命令解释"],
    ], [4350, 5010])
    code_block(doc, """
/* Core/Src/main.c or interrupt glue: keep it thin. */
void HAL_ADC_ConvCpltCallback(ADC_HandleTypeDef *hadc)
{
    if (hadc == &hadc1)
    {
        BSP_ADC_OnFullTransferFromISR();
    }
}
""", "代码示例 1  生成层只转交事件")

    heading(doc, "4.2 BSP：把板卡硬件包装成能力", 2)
    p(doc, "`BSP/` 代表 Board Support Package。它可以认识 `hadc1`、`huart1`、DMA stream 和 GPIO 引脚，但对上层说的是“启动采样”“发送一段数据”“点亮状态灯”。")
    table(doc, ["BSP 模块", "封装资源", "面向上层的能力"], [
        ["`bsp_adc_dma`", "ADC + DMA 缓冲", "启动/停止采集、报告半块/整块完成"],
        ["`bsp_sample_clock`", "TIM 触发源", "设置采样率、启停触发"],
        ["`bsp_transport_uart`", "USART + DMA/IRQ", "异步发送、接收数据块"],
        ["`bsp_storage`", "内部 Flash / EEPROM", "读写配置记录"],
        ["`bsp_board_io`", "LED、电源、告警引脚", "状态指示和板级控制"],
    ], [2200, 2850, 4310])
    code_block(doc, """
typedef void (*BSP_ADC_BlockCallback)(const uint16_t *data, uint32_t count);

void BSP_ADC_Init(BSP_ADC_BlockCallback on_block_ready);
int  BSP_ADC_Start(uint16_t *dma_buffer, uint32_t sample_count);
void BSP_ADC_Stop(void);

int  BSP_SAMPLE_CLOCK_SetRate(uint32_t sample_rate_hz);
void BSP_SAMPLE_CLOCK_Start(void);
void BSP_SAMPLE_CLOCK_Stop(void);
""", "代码示例 2  BSP 对采集服务暴露稳定能力")

    heading(doc, "4.3 Services：沉淀可复用的产品能力", 2)
    p(doc, "`Services/` 是平台化价值最高的一层。它实现采集引擎、数据块缓冲、帧协议、配置校验、标定和诊断；这些能力应尽可能在更换硬件后继续使用。")
    table(doc, ["服务模块", "职责", "禁止依赖"], [
        ["Acquisition", "采样启停、块队列、溢出统计、通道排列", "`hadc1`、GPIO 宏"],
        ["Protocol", "命令解析、数据帧编码、CRC、版本兼容", "UART/USB 具体发送"],
        ["Configuration", "参数默认值、校验、加载和持久化模型", "直接 Flash HAL 擦写"],
        ["Calibration", "增益/零偏修正、标定数据应用", "某个 ADC 寄存器配置"],
        ["Diagnostics", "丢块、异常、健康数据和错误事件", "LED 引脚操作"],
    ], [1850, 4100, 3410])

    heading(doc, "4.4 Application：描述产品行为", 2)
    p(doc, "`Application/` 面向用户需求组织设备工作：收到开始命令后如何进入 STREAMING 状态，停止时如何安全排空传输，异常时是否进入保护模式。它不应该“亲手”启动 ADC 或拼装底层数据。")
    code_block(doc, """
switch (command.type)
{
case CMD_START_CAPTURE:
    if (ACQ_Start(&command.capture_config) == ACQ_OK)
    {
        APP_SetState(APP_STATE_STREAMING);
    }
    break;

case CMD_STOP_CAPTURE:
    ACQ_Stop();
    APP_SetState(APP_STATE_IDLE);
    break;
}
""", "代码示例 3  应用层只表达设备行为")

    heading(doc, "4.5 Platform：一块具体板卡的装配说明", 2)
    p(doc, "`Platform/` 不是“万能业务层”，而是装配层。它决定这块板上启用哪个 BSP transport、以什么顺序初始化服务、硬件事件如何投递到服务。")
    code_block(doc, """
void PLATFORM_Init(void)
{
    BSP_BOARD_IO_Init();
    BSP_TRANSPORT_Init(PLATFORM_OnRxBytes, PLATFORM_OnTxDone);
    BSP_ADC_Init(PLATFORM_OnSampleBlockReady);
    CONFIG_Init(&board_storage_backend);
    ACQ_Init(&board_acq_backend);
    DIAG_Init();
}

void PLATFORM_OnSampleBlockReady(const uint16_t *samples, uint32_t count)
{
    ACQ_OnBlockReadyFromISR(samples, count);
}
""", "代码示例 4  Platform 完成板级依赖装配")

    heading(doc, "5. 建议目录与命名规范", 1)
    p(doc, "当前工程已经具备一级目录。进入真实采集开发后，建议在 `Services/` 和 `BSP/` 中继续按职责拆分，而不是继续扩张一个 `daq_service.c`。")
    code_block(doc, """
Application/
  Inc/ app.h app_state.h app_commands.h
  Src/ app.c acquisition_task.c communication_task.c supervision_task.c
BSP/
  Inc/ bsp_adc_dma.h bsp_sample_clock.h bsp_transport.h bsp_board_io.h
  Src/ bsp_adc_dma.c bsp_sample_clock.c bsp_transport_uart.c bsp_board_io.c
Services/
  Acquisition/ Inc/ acquisition_service.h sample_block.h
               Src/ acquisition_service.c sample_queue.c
  Protocol/    Inc/ frame_protocol.h command_protocol.h
               Src/ frame_protocol.c command_protocol.c crc16.c
  Configuration/ Inc/ device_config.h
                 Src/ device_config.c
  Diagnostics/ Inc/ diagnostics.h
               Src/ diagnostics.c
Platform/
  Inc/ platform.h platform_config.h
  Src/ platform.c platform_callbacks.c
Core/ Drivers/ Middlewares/ MDK-ARM/ Docs/
""", "图 3  推荐的中期目录结构")
    heading(doc, "5.1 命名规则", 2)
    table(doc, ["对象类型", "推荐命名", "示例"], [
        ["BSP 模块 API", "`BSP_<DEVICE>_<Verb>`", "`BSP_ADC_Start()`"],
        ["服务模块 API", "`<SERVICE>_<Verb>`", "`ACQ_GetBlock()`"],
        ["应用模块 API", "`APP_<Verb>`", "`APP_HandleCommand()`"],
        ["平台回调", "`PLATFORM_On<Event>`", "`PLATFORM_OnSampleBlockReady()`"],
        ["状态枚举", "`<MODULE>_STATE_<Name>`", "`ACQ_STATE_RUNNING`"],
        ["配置结构", "`<MODULE>_Config`", "`ACQ_Config`"],
    ], [1800, 3180, 4380])
    bullet(doc, "头文件只公开上层需要的类型与接口，私有缓冲、HAL 句柄和内部状态留在 `.c` 中。")
    bullet(doc, "避免使用 `common.c`、`utils.c` 作为无边界容器；工具函数也应归属明确职责。")
    bullet(doc, "模块名使用能力语言，不使用具体客户需求或临时调试名称。")

    heading(doc, "6. 数据采集链路如何平台化", 1)
    p(doc, "信息采集板卡的主链路应围绕“固定大小的样本块”设计，而不是围绕一次 HAL 回调设计。样本块是硬件采集与协议发送之间的稳定交接单元。")
    code_block(doc, """
Timer Trigger -> ADC Conversion -> DMA Circular Buffer
       |                                |
       |                                v
       |                        BSP_ADC block event
       |                                |
       v                                v
 sample rate                    Acquisition Service
 configuration                  queue / timestamp / overflow
                                        |
                                        v
                                Protocol Encoder
                                        |
                                        v
                                BSP_Transport -> Host
""", "图 4  上行数据链路")
    heading(doc, "6.1 推荐的数据块模型", 2)
    code_block(doc, """
typedef struct
{
    uint32_t sequence;
    uint32_t timestamp_us;
    uint32_t sample_rate_hz;
    uint16_t channel_mask;
    uint16_t sample_count;
    const uint16_t *samples;
} SampleBlock;
""", "代码示例 5  服务层使用的采样块结构")
    table(doc, ["字段", "为何属于服务层数据模型"], [
        ["`sequence`", "检测上位机丢帧或采集块跳变"],
        ["`timestamp_us`", "支持多源对齐、离线复盘与采样间隔分析"],
        ["`sample_rate_hz`", "让接收方知道当前数据意义"],
        ["`channel_mask`", "通道配置改变时协议仍可兼容"],
        ["`samples`", "承载原始数据，不绑定 DMA 实际缓冲实现"],
    ], [2100, 7260])
    heading(doc, "6.2 双缓冲与队列边界", 2)
    number(doc, "BSP 用 DMA 循环缓冲接收数据，并在半传输/全传输事件中识别刚完成的数据半块。")
    number(doc, "ISR 只把“已完成块索引”投递给采集服务或 RTOS 队列，立即返回。")
    number(doc, "采集任务取得块后添加序号、时间戳和统计信息，再交给协议层编码。")
    number(doc, "通信任务负责发送；若发送跟不上，诊断模块记录拥塞或丢块，而不是阻塞 ADC 回调。")
    callout(doc, "设计目标", "采样时序由硬件和 DMA 保证；CPU 负责搬运已完成的数据块与管理状态，不把不可预测的通信耗时带入采样路径。", GREEN)

    heading(doc, "7. 中断、回调与 RTOS 任务边界", 1)
    p(doc, "平台化必须同时解决实时性问题。最简单而可靠的规则是：ISR 只做常数时间工作，复杂逻辑都转入任务上下文。")
    table(doc, ["执行位置", "允许动作", "不应执行"], [
        ["ADC/DMA ISR", "确认来源、投递块事件、递增轻量计数", "CRC、滤波、阻塞发送、日志格式化"],
        ["UART/USB RX ISR", "收字节/块、投递接收事件", "命令解析、Flash 参数保存"],
        ["Acquisition Task", "管理采样块、缓冲与统计", "直接修改外设寄存器"],
        ["Communication Task", "解析命令、编码/发送帧", "控制 DMA 缓冲所有权"],
        ["Supervision Task", "健康检查、LED、看门狗、恢复", "执行高频数据搬运"],
    ], [1880, 3860, 3620])
    heading(doc, "7.1 ISR 到任务的事件接口", 2)
    code_block(doc, """
/* Called in interrupt context. Never block here. */
void ACQ_OnBlockReadyFromISR(uint8_t completed_half)
{
    BaseType_t wake = pdFALSE;
    xQueueSendFromISR(acq_block_queue, &completed_half, &wake);
    portYIELD_FROM_ISR(wake);
}

/* Called in acquisition task context. */
void ACQ_Process(void)
{
    uint8_t half;
    if (xQueueReceive(acq_block_queue, &half, 0) == pdPASS)
    {
        ACQ_PublishCompletedBlock(half);
    }
}
""", "代码示例 6  ISR 只交付事件，任务执行处理")
    heading(doc, "7.2 如何判断任务是否需要拆分", 2)
    bullet(doc, "采集速率较高或发送有拥塞风险时，采集任务与通信任务必须分开。")
    bullet(doc, "存在命令解析、参数落盘或固件升级时，通信控制路径不可与高速上行路径混在一个循环。")
    bullet(doc, "需要故障恢复和看门狗时，监督任务应独立于数据通道，避免数据堵塞使健康检测失效。")

    heading(doc, "8. 协议、配置、标定和诊断的归属", 1)
    heading(doc, "8.1 数据帧协议属于 Services/Protocol", 2)
    p(doc, "传输介质可能从 UART 改为 USB CDC、以太网或 CAN，但采样数据的语义不应随之变化。帧结构、CRC、序号和命令格式应由协议服务维护。")
    table(doc, ["帧字段", "推荐内容", "用途"], [
        ["同步头", "固定魔数 + 协议版本", "发现帧边界并支持升级"],
        ["类型/长度", "DATA / STATUS / RESPONSE", "区分消息并校验长度"],
        ["序号/时间戳", "递增序号 + 采样时刻", "丢帧检测和时间重建"],
        ["采集配置摘要", "采样率、通道掩码、量程", "解释载荷"],
        ["负载", "原始或处理后采样值", "核心数据"],
        ["校验", "CRC16/CRC32", "链路完整性检测"],
    ], [1880, 3670, 3810])
    code_block(doc, """
size_t FRAME_EncodeSampleBlock(const SampleBlock *block,
                               uint8_t *frame,
                               size_t capacity);

CMD_Result COMMAND_Parse(const uint8_t *data,
                         size_t length,
                         DeviceCommand *command);
""", "代码示例 7  协议服务不认识 UART 或 USB")

    heading(doc, "8.2 参数与标定属于 Services/Configuration 和 Calibration", 2)
    table(doc, ["内容", "服务职责", "BSP 职责"], [
        ["采样率、通道开关", "数据结构、合法性校验、版本管理", "将配置落实到 ADC/Timer"],
        ["校准系数", "应用零偏/增益修正和版本兼容", "读取/写入非易失存储"],
        ["通信参数", "选择策略和默认值", "配置实际串口或接口"],
        ["恢复出厂设置", "决定重置哪些逻辑参数", "执行擦除或写入动作"],
    ], [2000, 4300, 3060])

    heading(doc, "8.3 诊断属于 Services/Diagnostics", 2)
    bullet(doc, "采集块产生数量、发送数量、丢弃数量、DMA 溢出次数。")
    bullet(doc, "协议 CRC 错误、未知命令、超时、发送失败次数。")
    bullet(doc, "当前应用状态、采样率、通道配置、最近故障码。")
    bullet(doc, "启动原因、看门狗复位原因、校准数据有效性。")
    callout(doc, "经验规则", "状态灯是诊断结果的显示方式，不是诊断本身。因此 LED 操作在 BSP，错误判定和计数在 Diagnostics。", LIGHT_BLUE)

    heading(doc, "9. 典型功能开发时如何分配代码", 1)
    heading(doc, "9.1 增加 MCU 内部 ADC + DMA 连续采集", 2)
    table(doc, ["步骤", "所在层", "产出"], [
        ["CubeMX 配置 ADC、DMA、TIM TRGO", "Core", "初始化与 IRQ 回调"],
        ["封装启停与缓冲事件", "BSP", "`bsp_adc_dma` / `bsp_sample_clock`"],
        ["消费完成块并统计溢出", "Services", "`acquisition_service`"],
        ["编码 DATA 帧", "Services", "`frame_protocol`"],
        ["收到 START/STOP 后控制服务", "Application", "设备运行状态机"],
        ["组合初始化和事件回调", "Platform", "本板连接关系"],
    ], [1050, 1700, 6610])
    heading(doc, "9.2 从内部 ADC 切换为外部 SPI ADC", 2)
    p(doc, "如果前一阶段分层正确，切换外部 ADC 时，主要变化应发生在 CubeMX 与 BSP：SPI、DRDY 中断、DMA 和芯片寄存器配置会变，但 `SampleBlock`、数据帧协议、命令接口和应用状态机仍可复用。")
    table(doc, ["模块", "是否通常需要修改", "原因"], [
        ["Core", "需要", "配置 SPI/DMA/EXTI 与引脚"],
        ["BSP ADC", "需要重写或替换", "具体器件驱动发生变化"],
        ["Acquisition Service", "少量或无需", "仍接收标准数据块"],
        ["Protocol Service", "无需或仅增元数据", "上行帧语义保持稳定"],
        ["Application", "无需", "仍是启动、停止、状态控制"],
    ], [2200, 2300, 4860])
    heading(doc, "9.3 UART 切换 USB 或以太网", 2)
    bullet(doc, "创建新的 `bsp_transport_usb.c` 或 `bsp_transport_eth.c`，实现统一发送/接收接口。")
    bullet(doc, "Platform 根据板型或构建配置选择 transport backend。")
    bullet(doc, "Protocol 继续输出相同帧格式；上位机只需换物理连接或链路封装。")
    bullet(doc, "Application 不应感知底层究竟是 USART、USB 还是网络。")

    heading(doc, "10. API 设计与模块协作规则", 1)
    p(doc, "平台化最终是否稳定，很大程度取决于接口。好的接口表达业务语义、明确所有权、能报告失败，并避免把底层实现细节泄漏给上层。")
    table(doc, ["规则", "推荐做法", "避免做法"], [
        ["表达能力", "`ACQ_Start(&config)`", "`HAL_ADC_Start_DMA(&hadc1, ...)` 出现在应用层"],
        ["错误返回", "返回枚举状态并记录诊断", "静默失败或只在串口打印"],
        ["缓冲所有权", "说明谁创建、谁释放、有效期多久", "多任务随意共享 DMA 指针"],
        ["回调上下文", "接口名称标明 `FromISR`", "调用者猜测是否可阻塞"],
        ["配置版本", "结构体携带版本/长度", "参数以散乱全局变量保存"],
    ], [1720, 3770, 3870])
    heading(doc, "10.1 推荐服务接口草案", 2)
    code_block(doc, """
typedef enum
{
    ACQ_OK = 0,
    ACQ_ERR_CONFIG,
    ACQ_ERR_BUSY,
    ACQ_ERR_OVERFLOW,
    ACQ_ERR_HARDWARE
} ACQ_Result;

typedef struct
{
    uint32_t sample_rate_hz;
    uint16_t channel_mask;
    uint16_t samples_per_block;
} ACQ_Config;

void       ACQ_Init(const ACQ_Backend *backend);
ACQ_Result ACQ_Start(const ACQ_Config *config);
void       ACQ_Stop(void);
int        ACQ_TryGetBlock(SampleBlock *block);
ACQ_State  ACQ_GetState(void);
""", "代码示例 8  业务稳定、硬件可替换的采集接口")
    heading(doc, "10.2 配置注入而不是跨层全局访问", 2)
    p(doc, "服务层若需要使用板级能力，可由 Platform 在初始化时注入一组函数指针或 backend 结构。这样同一套服务可接入真实 BSP，也可以在 PC 测试中接入模拟 backend。")
    code_block(doc, """
typedef struct
{
    int  (*start)(uint16_t *buffer, uint32_t count);
    void (*stop)(void);
    int  (*set_rate)(uint32_t hz);
    uint32_t (*timestamp_us)(void);
} ACQ_Backend;
""", "代码示例 9  后端注入模型")

    heading(doc, "11. CubeMX 与 Keil 工程协作方式", 1)
    p(doc, "平台化并不排斥 CubeMX，而是规定生成代码的边界。只要手写功能层不被生成器拥有，后续增加外设和重新生成就不会反复破坏架构。")
    heading(doc, "11.1 每次新增外设的标准流程", 2)
    number(doc, "在 `.ioc` 中配置外设、DMA、中断和引脚，只让 CubeMX 生成底层初始化。")
    number(doc, "检查 `Core/` 中的新增文件和 `USER CODE` 回调接入点，不在自动区域写业务。")
    number(doc, "在 `BSP/` 创建适配模块，以能力接口封装 HAL 句柄。")
    number(doc, "在 `Services/` 添加或扩展服务状态、数据模型与错误处理。")
    number(doc, "在 `Platform/` 连接 BSP 与服务，并在 `Application/` 组织可见行为。")
    number(doc, "将新源码和 include 路径纳入 Keil 分组，执行完整 Rebuild 和板级验证。")
    heading(doc, "11.2 Keil 分组建议", 2)
    table(doc, ["Keil Group", "应包含"], [
        ["Application", "`app.c`、各任务和产品状态机"],
        ["Platform", "`platform.c`、回调装配和板型配置"],
        ["Services/Acquisition", "采集服务与队列"],
        ["Services/Protocol", "帧编码、命令解析、CRC"],
        ["Services/Diagnostics", "计数、故障和状态输出"],
        ["BSP", "板级硬件适配模块"],
        ["Application/User/Core", "CubeMX 生成的业务入口 glue"],
    ], [2700, 6660])
    callout(doc, "重新生成后的固定检查", "确认 `Core/Src/freertos.c` 中 `APP_Init()` / `APP_RunOnce()` 接入仍存在；确认手工创建的分层源码仍在 Keil 工程中；确认完整重建无警告。", YELLOW)

    heading(doc, "12. 测试、联调与质量控制", 1)
    p(doc, "平台化的收益必须通过可验证性兑现。采集板卡测试应同时覆盖硬件时序、数据正确性、协议完整性和异常恢复。")
    table(doc, ["测试层级", "重点检查", "建议方法"], [
        ["模块测试", "CRC、帧编码、配置校验、队列逻辑", "PC 仿真 backend 或小型单元测试"],
        ["BSP 联调", "ADC/DMA/Timer 是否按配置运行", "示波器、逻辑分析仪、固定输入"],
        ["系统联调", "上位机接收、序号连续、时间戳合理", "长时间流传输记录"],
        ["压力测试", "最高采样率、通信拥塞、缓冲溢出", "故意降速接收并检查统计"],
        ["故障测试", "断链、参数错误、复位恢复、看门狗", "故障注入与状态查询"],
    ], [1720, 3850, 3790])
    heading(doc, "12.1 每个数据路径都要有可观测量", 2)
    table(doc, ["指标", "用于发现的问题"], [
        ["`sample_blocks_generated`", "ADC/DMA 是否持续产出"],
        ["`sample_blocks_transmitted`", "通信是否跟上采集"],
        ["`sample_blocks_dropped`", "缓冲容量或带宽不足"],
        ["`protocol_crc_errors`", "链路或协议解析异常"],
        ["`last_fault_code`", "系统最近为何降级或停止"],
        ["`uptime_ms` / reset reason", "复位与稳定性问题"],
    ], [3100, 6260])
    heading(doc, "12.2 代码评审门槛", 2)
    bullet(doc, "应用层是否直接出现 HAL 句柄、DMA stream 或 GPIO pin？若出现，优先重构到 BSP。")
    bullet(doc, "ISR 是否执行了可能阻塞或耗时不确定的动作？")
    bullet(doc, "数据缓冲所有权和生命期是否能被清楚描述？")
    bullet(doc, "新增功能是否有错误返回、诊断指标和超时策略？")
    bullet(doc, "更换 transport 或 ADC 时，Services 和 Application 是否仍可复用？")

    heading(doc, "13. 从当前骨架到完整板卡平台的实施路线", 1)
    p(doc, "下面的路线以风险和依赖顺序组织。不要急于一次写全所有目录；每个阶段都应产生可验证结果，并保持层间边界成立。")
    table(doc, ["阶段", "核心任务", "可验收结果"], [
        ["阶段 0：骨架", "保持当前 Application/Platform/Services/BSP 接线", "现有工程构建通过"],
        ["阶段 1：采样点亮", "配置 ADC + Timer + DMA；建立 BSP ADC", "固定输入可产生稳定 DMA 块"],
        ["阶段 2：采集服务", "双缓冲、块序号、丢块统计", "任务侧持续获取 SampleBlock"],
        ["阶段 3：数据上行", "协议编码与 transport 异步发送", "上位机解帧并验证连续性"],
        ["阶段 4：控制面", "命令、参数、状态查询、持久化", "可在线配置采样并回读状态"],
        ["阶段 5：可靠性", "诊断、标定、看门狗、故障恢复", "长时间运行及故障注入通过"],
    ], [1800, 4200, 3360])
    heading(doc, "13.1 第一轮实现应创建的模块", 2)
    table(doc, ["优先级", "新增文件", "先做到什么程度"], [
        ["P0", "`BSP/bsp_adc_dma.*`", "启动 DMA，投递半块/整块事件"],
        ["P0", "`BSP/bsp_sample_clock.*`", "设置固定采样率并启停触发"],
        ["P0", "`Services/Acquisition/*`", "队列取块、状态与丢块计数"],
        ["P1", "`Services/Protocol/*`", "最小 DATA 帧和 CRC"],
        ["P1", "`BSP/bsp_transport_uart.*`", "异步或队列化发送"],
        ["P2", "`Application/*task.c`", "采集与通信任务解耦"],
    ], [1100, 3550, 4710])
    callout(doc, "不要提前复杂化", "滤波、压缩、升级和多通信介质可以后加。第一目标是建立一条边界正确、可测量丢块、可被上位机验证的数据采集通路。", GREEN)

    heading(doc, "附录 A：模块职责速查表", 1)
    table(doc, ["我正在写的代码...", "它应放在哪里", "理由"], [
        ["设置 ADC 通道、DMA 触发、中断优先级", "Core / CubeMX", "芯片初始化事实"],
        ["调用 `HAL_ADC_Start_DMA`", "BSP", "具体硬件动作"],
        ["决定采样率是否合法", "Services/Configuration", "产品规则"],
        ["管理采样块队列与丢块", "Services/Acquisition", "可复用能力"],
        ["拼装帧头、CRC 与命令", "Services/Protocol", "链路无关语义"],
        ["执行 UART/USB 发送", "BSP/Transport", "介质实现"],
        ["收到开始命令后切换运行模式", "Application", "产品行为"],
        ["选择 UART backend 并连接回调", "Platform", "板级装配"],
        ["点亮错误 LED", "BSP/Board IO", "硬件显示动作"],
        ["判断何时显示错误 LED", "Diagnostics/Application", "策略和状态"],
    ], [4020, 2500, 2840])

    heading(doc, "附录 B：开发检查清单", 1)
    heading(doc, "B.1 新增硬件能力前", 2)
    bullet(doc, "明确该能力属于采集、通信、存储、电源还是诊断。")
    bullet(doc, "确认 CubeMX 配置会新增哪些 HAL 句柄和回调入口。")
    bullet(doc, "先定义 BSP 对上层需要提供的最小接口，再实现 HAL 调用。")
    bullet(doc, "明确中断上下文和任务上下文之间的数据交付方式。")
    heading(doc, "B.2 新增服务功能前", 2)
    bullet(doc, "服务输入输出是否独立于具体硬件实例。")
    bullet(doc, "状态、错误码、统计指标是否完整。")
    bullet(doc, "缓冲大小、所有权和超载策略是否写明。")
    bullet(doc, "是否能借助模拟 backend 进行逻辑测试。")
    heading(doc, "B.3 合并与发布前", 2)
    bullet(doc, "CubeMX 重新生成后平台接入代码仍保留。")
    bullet(doc, "Keil 完整 Rebuild 为 0 error / 0 warning。")
    bullet(doc, "目标采样率下无不可解释丢块，故障计数可读取。")
    bullet(doc, "协议文档、配置默认值和板级引脚说明已更新。")
    bullet(doc, "烧录镜像、版本号和测试记录能够对应。")

    heading(doc, "附录 C：术语表", 1)
    table(doc, ["术语", "含义"], [
        ["Platformization / 平台化", "通过稳定边界和可替换实现，使同一套业务能力适配多种硬件或产品配置。"],
        ["BSP", "Board Support Package，封装一块板上的具体硬件操作。"],
        ["Service", "独立于具体板卡实现的功能能力，如采集、协议、诊断。"],
        ["Backend", "供服务调用的可替换底层能力集合，通常由 BSP 提供。"],
        ["SampleBlock", "从采样硬件交给上层处理的固定语义数据块。"],
        ["Transport", "承载命令或数据帧的通信介质，如 UART、USB、Ethernet。"],
        ["ISR", "中断服务例程，要求执行快速、确定且不可阻塞。"],
        ["Glue Code", "位于生成层与手写层边界上的最小连接代码。"],
    ], [2600, 6760])
    heading(doc, "结语", 1)
    p(doc, "平台化不是为了让目录看起来整齐，而是为了让固件在硬件、需求和时间的变化下仍然可以被理解、替换和验证。对 DAQ_card 来说，最重要的下一步不是立刻堆功能，而是沿着本文的边界建立第一条真实采集通路：Core 配置硬件，BSP 提供能力，Services 稳定数据语义，Application 控制设备行为，Platform 把这一切装配成一块可持续演进的板卡。")
    callout(doc, "落地起点", "下一次开发工作应从 `ADC + Timer Trigger + DMA Circular` 的 CubeMX 配置开始，同时创建 `bsp_adc_dma` 与 `acquisition_service`，让第一批真实样本沿平台链路抵达上位机。", LIGHT_BLUE)


def main():
    doc = Document()
    configure_document(doc)
    configure_numbering(doc)
    cover(doc)
    content_pages(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "信息采集板卡固件平台化开发教程"
    doc.core_properties.subject = "STM32 数据采集固件平台化架构与落地方法"
    doc.core_properties.author = "DAQ_card Firmware Team"
    doc.core_properties.keywords = "STM32, DAQ, Firmware, Platformization, BSP, FreeRTOS"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
