from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
BASE_PATH = ROOT / "Docs" / "source" / "build_platform_tutorial.py"
OUTPUT = ROOT / "Docs" / "DAQ_Card_GitHub使用与密钥安全指南.docx"

spec = spec_from_file_location("doc_base", BASE_PATH)
base = module_from_spec(spec)
spec.loader.exec_module(base)


def set_running_header(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("DAQ_CARD  |  GITHUB WORKFLOW AND KEY SECURITY GUIDE")
    base.set_font(run, 8.5, base.MUTED, bold=True)
    base.paragraph_border_bottom(header, base.BORDER, "6", "5")

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("GitHub 使用与密钥安全指南  |  ")
    base.set_font(run, 8.5, base.MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def cover(doc):
    kicker = doc.add_paragraph("VERSION CONTROL HANDBOOK", style="Subtitle")
    kicker.runs[0].font.color.rgb = RGBColor.from_string(base.BLUE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(45)
    title = doc.add_paragraph("GitHub 使用与\n密钥安全指南", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph("面向 DAQ_Card 固件工程的日常操作、推送与故障排查", style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rule = doc.add_paragraph()
    base.paragraph_border_bottom(rule, base.BLUE, "14", "12")
    base.callout(
        doc,
        "先回答你的问题",
        "上次推送使用的是一个临时 SSH deploy key：它只临时授权本仓库写入，完成推送后已从 GitHub 删除，本地私钥也已删除。它不是你的 GitHub 密码，也没有作为文件提交到仓库。",
        base.LIGHT_BLUE,
    )
    base.table(doc, ["项目", "说明"], [
        ["目标仓库", "https://github.com/emptyzcb/DAQ_Card"],
        ["默认分支", "main"],
        ["当前首个提交", "31dbde3 - Initial platformized DAQ card firmware project"],
        ["适用读者", "需要维护、提交和发布本固件工程的开发人员"],
        ["文档版本", "V1.0  |  2026-05-27"],
    ], [2000, 7360])
    doc.add_page_break()


def body(doc):
    base.heading(doc, "阅读路线", 1)
    base.p(doc, "这份指南先把“密钥到底是什么”讲清楚，再带你建立可靠的 GitHub 使用习惯。所有命令均以本项目 `D:\\A_stm32_project\\DAQ_card` 和仓库 `emptyzcb/DAQ_Card` 为例。")
    base.table(doc, ["章节", "用途"], [
        ["1-2", "理解临时密钥、Git、GitHub、远端和分支"],
        ["3-5", "掌握认证方式、日常提交推送和分支协作"],
        ["6-8", "管理 STM32 工程文件、文档、发布物与安全"],
        ["9-10", "解决推送失败、回退错误和形成固定习惯"],
        ["附录", "命令速查与官方资料链接"],
    ], [1500, 7860])

    base.heading(doc, "1. 刚才使用的密钥是什么", 1)
    base.p(doc, "在将工程首次推送到 GitHub 时，普通 HTTPS Git 连接在当前网络上无法连接到 `github.com:443`，而 GitHub 官方提供的 SSH-over-443 入口可以连接。为了完成一次推送，使用了一个临时的 SSH deploy key。")
    base.heading(doc, "1.1 临时 deploy key 的完整过程", 2)
    base.number(doc, "本机临时生成一对 ED25519 SSH 密钥：公钥和私钥。")
    base.number(doc, "仅将公钥注册到 `emptyzcb/DAQ_Card` 仓库，并临时允许写入。")
    base.number(doc, "Git 使用私钥，通过 `ssh.github.com:443` 将 `main` 分支推送到该仓库。")
    base.number(doc, "推送成功后，从 GitHub 仓库删除公钥，并删除本机临时私钥文件。")
    base.number(doc, "最终复核仓库 deploy key 列表为空，说明这个临时通道已失效。")
    base.table(doc, ["要素", "在哪里", "现在的状态"], [
        ["公钥", "曾短暂登记在该 GitHub 仓库设置中", "已删除"],
        ["私钥", "只存在于本机临时目录，用于签名 SSH 连接", "已删除"],
        ["授权范围", "仅该仓库的一次写入用途", "已失效"],
        ["提交内容", "源码、工程文件和文档", "不包含任何密钥"],
    ], [1750, 4580, 3030])
    base.callout(
        doc,
        "结论",
        "你不需要保管这把临时钥匙，也不需要撤销任何遗留权限。以后若需要稳定地通过 SSH 推送，应该创建属于你账号的长期 SSH key，并将公钥添加到账号，而不是复用临时 deploy key。",
        base.GREEN,
    )

    base.heading(doc, "1.2 三种容易混淆的认证对象", 2)
    base.table(doc, ["对象", "它是什么", "本项目中的情况"], [
        ["GitHub CLI 登录令牌", "`gh auth login` 后安全保存的账号认证凭据，用于调用 GitHub API", "本机已有 `emptyzcb` 登录状态；不要打印或提交 token"],
        ["Deploy key", "绑定到单个仓库的 SSH 公钥；私钥留在部署机或自动化环境", "曾临时使用，现已删除"],
        ["SSH 主机指纹 / known_hosts", "用于确认连接的是 GitHub 服务器，而非冒充者", "首次连接 `ssh.github.com:443` 出现提示属于正常安全校验"],
    ], [2150, 4670, 2540])
    base.callout(
        doc,
        "安全记忆法",
        "账号 token 代表“我是谁”；deploy key 代表“这个自动化主体可以访问哪个仓库”；host key 代表“对方服务器真的是 GitHub”。",
        base.LIGHT_BLUE,
    )

    base.heading(doc, "2. Git 与 GitHub 的工作模型", 1)
    base.p(doc, "Git 是本机的版本控制工具，GitHub 是保存和协作共享 Git 提交的远端服务。即使断网，你也能在本地提交；只有 fetch、pull、push 和网页协作功能需要远端。")
    base.code_block(doc, """
工作区文件 -> git add -> 暂存区 -> git commit -> 本地 main
                                                  |
                                                  v
                                         git push origin main
                                                  |
                                                  v
                                      GitHub: origin/main
""", "图 1  日常提交与推送路径")
    base.table(doc, ["术语", "在本工程中的含义"], [
        ["Repository / 仓库", "包含代码、历史和配置的版本库，即 DAQ_Card"],
        ["Commit / 提交", "一次可追溯的变更快照，例如首次平台化提交 `31dbde3`"],
        ["Branch / 分支", "独立开发线；正式主线当前为 `main`"],
        ["Remote / 远端", "本地 Git 为 GitHub 地址起的名字，默认叫 `origin`"],
        ["Push / 推送", "将本地已有提交发送到 GitHub"],
        ["Pull / 拉取", "取得远端更新并合并到当前本地分支"],
    ], [2450, 6910])
    base.heading(doc, "2.1 当前 DAQ_Card 仓库状态", 2)
    base.code_block(doc, """
Remote:  origin -> https://github.com/emptyzcb/DAQ_Card.git
Branch:  main tracks origin/main
Commit:  31dbde3 Initial platformized DAQ card firmware project
Docs:    Docs/信息采集板卡固件平台化开发教程.docx
Ignored: MDK-ARM/DAQ_card/ build outputs and Keil user-state files
""", "代码块 1  推送完成后的工程基线")

    base.heading(doc, "3. 推荐的认证方式", 1)
    base.heading(doc, "3.1 常规首选：GitHub CLI + HTTPS", 2)
    base.p(doc, "多数网络环境下，最简单的方式是使用 GitHub CLI 登录，再让 Git 使用 HTTPS 远端。登录凭据由系统凭据存储管理，不应写入源码目录。")
    base.code_block(doc, """
gh auth login
gh auth status
git remote -v
git push -u origin main
""", "代码块 2  常规认证和首次推送")
    base.bullet(doc, "执行 `gh auth status` 可确认当前活动账号；不要使用显示 token 的选项截图或分享。")
    base.bullet(doc, "远端地址为 HTTPS 时，通常无需自行管理 SSH 私钥。")
    base.bullet(doc, "任何密码、token、私钥文件都不允许加入 Git 提交。")
    base.heading(doc, "3.2 受限网络下：账号 SSH key + SSH 443", 2)
    base.p(doc, "本次网络可访问 GitHub API，却无法通过 HTTPS Git 连接 `github.com`。GitHub 官方支持通过 `ssh.github.com` 的 443 端口连接 SSH。若你今后在同一网络长期开发，可配置你本人账号的 SSH key。")
    base.code_block(doc, """
# 生成属于你账号的长期公私钥对（妥善保管私钥）
ssh-keygen -t ed25519 -C "your_email@example.com"

# 将 .pub 公钥内容添加到 GitHub 账号 Settings > SSH and GPG keys

# 测试官方 SSH-over-443 端点
ssh -T -p 443 git@ssh.github.com

# 本仓库切换 SSH 443 远端
git remote set-url origin ssh://git@ssh.github.com:443/emptyzcb/DAQ_Card.git
git push -u origin main
""", "代码块 3  持续使用 SSH 443 的正确做法")
    base.table(doc, ["选择", "适用场景", "密钥管理"], [
        ["HTTPS + `gh auth login`", "个人日常开发、网络正常", "由 GitHub CLI 和凭据存储管理 token"],
        ["账号 SSH key", "常用 SSH 或 HTTPS Git 被网络阻断", "私钥由本人长期保管，公钥绑定账号"],
        ["Deploy key", "服务器部署、CI 或单仓库自动化", "只绑定单个仓库，不适合替代个人日常登录"],
    ], [2700, 3580, 3080])

    base.heading(doc, "4. 个人日常开发流程", 1)
    base.p(doc, "固件项目推荐形成稳定的小步提交节奏：开始工作前同步；修改后先审查文件；构建通过后提交；需要共享时推送。")
    base.heading(doc, "4.1 每次开始工作", 2)
    base.code_block(doc, """
cd D:\\A_stm32_project\\DAQ_card
git status
git branch --show-current
git pull --ff-only origin main
""", "代码块 4  开工前检查")
    base.bullet(doc, "`git status` 应先确认没有遗忘的修改，避免把不同目的的改动混在一起。")
    base.bullet(doc, "`git pull --ff-only` 只接受线性同步，出现分叉时要求你明确处理，而不是自动制造混乱的合并。")
    base.heading(doc, "4.2 修改、检查与提交", 2)
    base.code_block(doc, """
# 查看变化
git status --short
git diff

# 只加入本次相关文件
git add Application/ Services/ BSP/ Platform/ Core/ MDK-ARM/DAQ_card.uvprojx
git add Docs/ README.md DAQ_card.ioc

# 检查即将提交的内容
git diff --cached --stat
git diff --cached

# 创建提交
git commit -m "Add ADC DMA acquisition service skeleton"
""", "代码块 5  形成一笔清晰提交")
    base.callout(doc, "固件项目的提交习惯", "提交前至少完成一次 Keil Build；凡是涉及采样时序、协议或板级适配的变更，还应记录板上验证结果或测试条件。", base.GREEN)
    base.heading(doc, "4.3 推送和验证", 2)
    base.code_block(doc, """
git push origin main
git status --short --branch
git log --oneline --decorate -3
""", "代码块 6  推送后确认本地与远端一致")
    base.p(doc, "看到 `main...origin/main` 且没有待提交变更，说明本地代码与 GitHub 主线同步。仅显示 `!!` 的 Keil 构建输出表示这些文件被 `.gitignore` 正常排除。")

    base.heading(doc, "5. 更稳健的分支与 Pull Request 工作流", 1)
    base.p(doc, "当功能开始变复杂，直接向 `main` 提交会降低可审查性。建议每项明显功能在独立分支中完成，通过 Pull Request 合入主线。")
    base.code_block(doc, """
git switch main
git pull --ff-only origin main
git switch -c feature/adc-dma-acquisition

# 修改、构建、提交
git add ...
git commit -m "Implement ADC DMA block capture"
git push -u origin feature/adc-dma-acquisition

# 创建 Pull Request
gh pr create --base main --head feature/adc-dma-acquisition
""", "代码块 7  功能分支工作流")
    base.table(doc, ["分支前缀", "用途", "示例"], [
        ["`feature/`", "新增可见能力", "`feature/adc-dma-acquisition`"],
        ["`fix/`", "修正缺陷", "`fix/uart-frame-overflow`"],
        ["`docs/`", "文档或规范更新", "`docs/github-guide`"],
        ["`refactor/`", "不改变功能的结构整理", "`refactor/transport-interface`"],
    ], [1800, 2900, 4660])
    base.heading(doc, "5.1 Pull Request 应写什么", 2)
    base.bullet(doc, "改了什么：模块、外设、接口和文档变化。")
    base.bullet(doc, "为什么改：对应板卡需求、问题或后续扩展目标。")
    base.bullet(doc, "如何验证：Keil 构建结果、硬件条件、采样率和观察结果。")
    base.bullet(doc, "风险是什么：中断时序、缓冲容量、兼容性或烧录影响。")

    base.heading(doc, "6. STM32 / Keil 工程该提交哪些文件", 1)
    base.p(doc, "嵌入式仓库既要能复现构建，又要避免把每次编译都变化的中间文件推上去。当前工程已经配置 `.gitignore` 和 `.gitattributes` 来处理这一点。")
    base.table(doc, ["文件类别", "是否提交", "本项目示例与理由"], [
        ["应用与平台源码", "提交", "`Application/`、`BSP/`、`Services/`、`Platform/`"],
        ["CubeMX 配置与生成源码", "提交", "`DAQ_card.ioc`、`Core/`，支持重生成与构建"],
        ["供应商依赖", "当前提交", "`Drivers/`、`Middlewares/`，保证工程离线可构建"],
        ["Keil 工程定义", "提交", "`MDK-ARM/DAQ_card.uvprojx`、startup 文件"],
        ["架构与 Word 文档", "提交", "`Docs/` 和 `README.md`"],
        ["编译结果", "不提交", "`.o/.axf/.map/.hex/.build_log.htm` 可重新生成"],
        ["个人 IDE 状态", "不提交", "`.uvoptx/.uvguix.*` 与用户环境相关"],
    ], [2330, 1530, 5500])
    base.heading(doc, "6.1 为什么 `.hex` 通常不直接纳入主线", 2)
    base.p(doc, "源码仓库应以可重现的源码和配置为准。固件镜像更适合在形成版本发布时作为 GitHub Release 附件上传，并同时记录提交哈希、构建工具版本、硬件版本和测试结论。这样开发历史不会被大量二进制产物污染，同时仍能交付可烧录文件。")
    base.heading(doc, "6.2 发布固件建议", 2)
    base.code_block(doc, """
# 当功能通过验证后创建标签
git tag -a v0.1.0 -m "First acquisition bring-up release"
git push origin v0.1.0

# 使用 GitHub CLI 创建发布并附加产物（示例）
gh release create v0.1.0 MDK-ARM/DAQ_card/DAQ_card.hex --title "DAQ Card v0.1.0"
""", "代码块 8  将烧录产物作为发布附件管理")

    base.heading(doc, "7. 密钥、Token 与敏感信息安全", 1)
    base.callout(doc, "绝对规则", "不要把密码、GitHub token、SSH 私钥、设备量产密钥、Wi-Fi 密码、客户数据或校准机密写入 Git 提交。仓库一旦公开推送，即使后续删除文件，也可能已经泄露。", base.RED)
    base.heading(doc, "7.1 哪些东西可以提交，哪些不能", 2)
    base.table(doc, ["对象", "可以提交吗", "处理方式"], [
        ["SSH 公钥 `.pub`", "谨慎，通常不需要", "添加到 GitHub 设置而不是工程源码"],
        ["SSH 私钥", "绝对不可", "只保存在用户密钥目录并限制访问权限"],
        ["`gh` / PAT token", "绝对不可", "保存在系统凭据存储或安全 CI Secret"],
        ["公开协议文档", "可以", "放入 `Docs/` 并随版本维护"],
        ["设备密钥/生产凭据", "绝对不可", "使用专门密钥管理和注入流程"],
        ["示例配置", "可以", "使用无效示例值，并明确标注 placeholder"],
    ], [2650, 1900, 4810])
    base.heading(doc, "7.2 提交前自检", 2)
    base.code_block(doc, """
git status --short
git diff --cached --name-only
git diff --cached

# 重点查看是否误加入常见敏感文件
git ls-files | findstr /i ".pem .key id_rsa id_ed25519 token secret password"
""", "代码块 9  提交前安全检查")
    base.heading(doc, "7.3 如果误提交了密钥", 2)
    base.number(doc, "第一时间在 GitHub 或对应系统撤销/轮换密钥或 token，不要只删除文件。")
    base.number(doc, "停止继续传播该提交，确认泄露范围。")
    base.number(doc, "视情况清理 Git 历史，并告知受影响协作者重新同步。")
    base.number(doc, "记录原因并增加 `.gitignore`、Secret 扫描或流程检查，避免再次发生。")

    base.heading(doc, "8. 常用查看、撤回与恢复操作", 1)
    base.table(doc, ["想做什么", "安全命令", "说明"], [
        ["查看当前修改", "`git status` / `git diff`", "不改变任何文件"],
        ["取消尚未暂存的单文件修改", "`git restore <file>`", "会丢掉该文件未提交修改，先确认"],
        ["取消暂存但保留修改", "`git restore --staged <file>`", "适合误 `git add`"],
        ["修改最近提交说明", "`git commit --amend`", "尚未推送时最安全"],
        ["撤回已推送提交", "`git revert <sha>`", "创建反向提交，保留公共历史"],
        ["查看某提交", "`git show <sha>`", "确认版本内容和变更原因"],
    ], [2350, 3050, 3960])
    base.callout(doc, "公共分支规则", "已经推送到 `main` 的提交，优先使用 `git revert` 纠错。不要随意强制推送或重写主线历史。", base.YELLOW)

    base.heading(doc, "9. 推送故障排查", 1)
    base.heading(doc, "9.1 连接 `github.com:443` 超时", 2)
    base.p(doc, "这正是首次推送时遇到的情况：Git 的 HTTPS 连接失败，但 GitHub API 和 `ssh.github.com:443` 可达。它通常是网络出口、防火墙或代理对不同目标/协议处理不同造成的。")
    base.table(doc, ["检查", "命令", "判断"], [
        ["GitHub 登录状态", "`gh auth status`", "确认账号凭据本身正常"],
        ["HTTPS Git 连通性", "`git ls-remote https://github.com/emptyzcb/DAQ_Card.git`", "超时说明 HTTPS Git 路径受阻"],
        ["SSH 443 连通性", "`ssh -T -p 443 git@ssh.github.com`", "可连接则可改用 SSH 路径"],
        ["远端 URL", "`git remote -v`", "确认正在使用的传输方案"],
    ], [2100, 4570, 2690])
    base.code_block(doc, """
# 使用个人 SSH key 后，将远端改为官方 SSH 443 形式
git remote set-url origin ssh://git@ssh.github.com:443/emptyzcb/DAQ_Card.git
git push origin main
""", "代码块 10  HTTPS 被阻断后的长期解决方向")
    base.heading(doc, "9.2 Push 被拒绝：远端已有新提交", 2)
    base.code_block(doc, """
git fetch origin
git status --branch
git pull --rebase origin main
# 解决冲突并验证后
git push origin main
""", "代码块 11  先整合远端更新再推送")
    base.heading(doc, "9.3 文件过大或仓库膨胀", 2)
    base.bullet(doc, "不要提交 Keil 构建目录和重复二进制输出；现有 `.gitignore` 已覆盖常见产物。")
    base.bullet(doc, "大型采集录波数据和测试数据应放在外部存储或发布附件中，不进入源码历史。")
    base.bullet(doc, "确实需要追踪的大型二进制资产，应评估 Git LFS，而不是直接反复提交。")

    base.heading(doc, "10. 面向本项目的固定工作规范", 1)
    base.table(doc, ["时机", "必须动作", "推荐结果"], [
        ["开始开发", "`git status`、同步 `main`、新建功能分支", "开发基线清晰"],
        ["修改 CubeMX 配置", "提交 `.ioc`、生成源码和 Keil 工程变更", "外设配置可追溯"],
        ["修改平台层", "更新相应接口说明或架构文档", "代码职责可理解"],
        ["完成功能", "Keil Rebuild、硬件验证、清晰 commit", "提交可复查"],
        ["准备合入", "PR 写明验证和风险", "主线质量可控"],
        ["准备交付", "创建 tag/release 并附 `.hex` 与测试说明", "固件可复现"],
    ], [1700, 4730, 2930])
    base.heading(doc, "10.1 提交信息建议", 2)
    base.code_block(doc, """
Add ADC DMA acquisition skeleton
Implement UART sample frame encoding
Fix DMA block overflow handling
Document GitHub workflow and key security
Release DAQ card firmware v0.1.0
""", "代码块 12  清晰的提交主题示例")
    base.heading(doc, "10.2 最小日常速记", 2)
    base.code_block(doc, """
git status
git pull --ff-only origin main
git switch -c feature/your-change

# edit + build + test
git add <related files>
git diff --cached
git commit -m "Describe the completed change"
git push -u origin feature/your-change
gh pr create --base main
""", "代码块 13  推荐的日常循环")

    base.heading(doc, "附录 A：命令速查", 1)
    base.table(doc, ["目的", "命令"], [
        ["查看登录账号", "`gh auth status`"],
        ["查看远端", "`git remote -v`"],
        ["查看工作区", "`git status --short --branch`"],
        ["查看修改", "`git diff` / `git diff --cached`"],
        ["查看历史", "`git log --oneline --decorate --graph -10`"],
        ["创建分支", "`git switch -c feature/name`"],
        ["提交", "`git add ...` 然后 `git commit -m \"Message\"`"],
        ["推送新分支", "`git push -u origin feature/name`"],
        ["同步主线", "`git pull --ff-only origin main`"],
        ["撤销公共提交", "`git revert <commit-sha>`"],
        ["SSH 443 测试", "`ssh -T -p 443 git@ssh.github.com`"],
    ], [2650, 6710])

    base.heading(doc, "附录 B：官方参考资料", 1)
    base.p(doc, "下列页面是本文关于认证、远端、deploy key 与 SSH 443 连接说明的官方依据；查阅日期：2026-05-27。")
    official = [
        ["添加本地代码到 GitHub", "https://docs.github.com/en/migrations/importing-source-code/using-the-command-line-to-import-source-code/adding-locally-hosted-code-to-github"],
        ["管理远端仓库", "https://docs.github.com/en/get-started/git-basics/managing-remote-repositories"],
        ["使用 SSH over HTTPS port", "https://docs.github.com/en/authentication/troubleshooting-ssh/using-ssh-over-the-https-port"],
        ["Deploy keys API 与概念", "https://docs.github.com/en/rest/deploy-keys/deploy-keys"],
        ["GitHub CLI 登录", "https://cli.github.com/manual/gh_auth_login"],
        ["GitHub CLI 登录状态", "https://cli.github.com/manual/gh_auth_status"],
    ]
    base.table(doc, ["主题", "官方链接"], official, [2450, 6910])

    base.heading(doc, "结语", 1)
    base.p(doc, "GitHub 的价值不只是“把代码放上网”，而是让每次硬件配置、平台接口、采集功能和发布固件都有来源、有验证、有回退路径。对 DAQ_Card 项目，最重要的习惯是：代码按职责提交，产物按发布交付，凭据始终留在安全存储中。")
    base.callout(
        doc,
        "从今天开始",
        "后续功能开发建议使用独立分支和 Pull Request；提交前检查敏感信息与构建结果；遇到当前网络的 HTTPS Git 超时问题时，使用你的个人 SSH key 配合 GitHub 官方 SSH 443 方案。",
        base.GREEN,
    )


def main():
    doc = Document()
    base.configure_document(doc)
    base.configure_numbering(doc)
    set_running_header(doc)
    cover(doc)
    body(doc)
    doc.core_properties.title = "DAQ_Card GitHub 使用与密钥安全指南"
    doc.core_properties.subject = "GitHub 操作、认证方式、Deploy Key 说明与固件仓库实践"
    doc.core_properties.author = "DAQ_Card Firmware Team"
    doc.core_properties.keywords = "GitHub, Git, SSH, Deploy Key, STM32, DAQ_Card"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
