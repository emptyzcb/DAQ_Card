from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "Docs" / "industrial_gyroscope_learning_path.md"
OUTPUT = ROOT / "Docs" / "industrial_gyroscope_learning_path.docx"

FONT_LATIN = "Calibri"
FONT_EAST_ASIA = "Microsoft YaHei"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "B7C9DD"


def set_run_font(run, size=None, bold=None, color=None, name=FONT_LATIN):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size, bold=False, color=None):
    font = style.font
    font.name = FONT_LATIN
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    font.size = Pt(size)
    font.bold = bold
    if color is not None:
        font.color.rgb = color


def set_paragraph_format(style, before=0, after=6, line=1.25):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().tcW
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def add_text_with_inline_code(paragraph, text, bold=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            set_run_font(run, size=9.5, color=DARK_BLUE, name="Consolas")
        else:
            set_run_font(run, size=11, bold=bold)


def next_id(elements, attr_name):
    values = []
    for element in elements:
        value = element.get(qn(attr_name))
        if value is not None and value.isdigit():
            values.append(int(value))
    return (max(values) + 1) if values else 1


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = next_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    num_id = next_id(numbering.findall(qn("w:num")), "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")

    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_heading(doc, text, level):
    style_name = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.keep_with_next = True
    add_text_with_inline_code(p, text, bold=True)
    return p


def add_code_block(doc, lines):
    for idx, line in enumerate(lines):
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(3 if idx == 0 else 0)
        p.paragraph_format.space_after = Pt(3 if idx == len(lines) - 1 else 0)
        run = p.add_run(line if line else " ")
        set_run_font(run, size=9.5, name="Consolas", color=DARK_BLUE)
        p_pr = p._p.get_or_add_pPr()
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), LIGHT_GRAY)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in rows[1]):
        return rows[:1] + rows[2:], i
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    set_table_borders(table)

    if col_count == 2:
        widths = [2700, 6660]
    elif col_count == 3:
        widths = [2200, 3580, 3580]
    else:
        widths = [9360 // col_count] * col_count
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            if r_idx == 0:
                shade_cell(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_text_with_inline_code(p, value, bold=(r_idx == 0))
            if c_idx == 0 and len(value) <= 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build_document():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11)
    set_paragraph_format(normal, 0, 6, 1.25)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        if name.endswith("1"):
            set_style_font(style, 16, True, BLUE)
            set_paragraph_format(style, 18, 10, 1.25)
        elif name.endswith("2"):
            set_style_font(style, 13, True, BLUE)
            set_paragraph_format(style, 14, 7, 1.25)
        else:
            set_style_font(style, 12, True, DARK_BLUE)
            set_paragraph_format(style, 10, 5, 1.25)

    title_style = styles.add_style("Learning Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title_style, 22, True, RGBColor(0x0B, 0x25, 0x45))
    set_paragraph_format(title_style, 0, 10, 1.15)

    subtitle_style = styles.add_style("Subtitle Small", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(subtitle_style, 10, False, MUTED)
    set_paragraph_format(subtitle_style, 0, 14, 1.15)

    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code_style, 9.5, False, DARK_BLUE)
    set_paragraph_format(code_style, 0, 0, 1.1)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        set_style_font(style, 11)
        set_paragraph_format(style, 0, 4, 1.25)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header.paragraphs[0]
    header.text = "工业化陀螺仪学习流程"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("第 ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    footer.runs[-1]._r.append(field_begin)
    footer.runs[-1]._r.append(instr)
    footer.runs[-1]._r.append(field_end)
    footer.add_run(" 页")
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    i = 0
    in_code = False
    code_lines = []
    first_heading = True
    current_num_id = None
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            current_num_id = None
            i += 1
            continue

        if line.startswith("|"):
            current_num_id = None
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and first_heading:
                p = doc.add_paragraph(style="Learning Title")
                add_text_with_inline_code(p, text, bold=True)
                meta = doc.add_paragraph(style="Subtitle Small")
                add_text_with_inline_code(meta, "面向 STM32 数据采集板卡项目：从基础概念到工业验证的阶段化路线")
                first_heading = False
            else:
                add_heading(doc, text, level)
            current_num_id = None
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
        bullet = re.match(r"^-\s+(.*)$", line.strip())
        if numbered:
            if current_num_id is None:
                current_num_id = create_decimal_numbering(doc)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            apply_numbering(p, current_num_id)
            add_text_with_inline_code(p, numbered.group(2))
        elif bullet:
            current_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_text_with_inline_code(p, bullet.group(1))
        else:
            current_num_id = None
            p = doc.add_paragraph()
            add_text_with_inline_code(p, line)
        i += 1

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
